#!/usr/bin/env python3
"""Literature Gap Radar.

This script intentionally starts from the user's local library before searching
new literature. It uses only the Python standard library for core operations;
PDF, DOCX, Excel, and Word support are best-effort when optional packages are
available.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import hashlib
import html
import json
import math
import os
import re
import sqlite3
import sys
import time
import urllib.parse
import urllib.request
import zipfile
from collections import Counter
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET


DEFAULT_WEIGHTS = {
    "topic_relevance": 0.28,
    "gap_fill": 0.20,
    "experiment_reusability": 0.16,
    "mechanism_value": 0.14,
    "novelty_against_local": 0.12,
    "journal_year_signal": 0.06,
    "evidence_richness": 0.04,
}

TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
FULLTEXT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
METADATA_EXTENSIONS = {".ris", ".bib", ".nbib", ".xml"}
DOC_EXTENSIONS = FULLTEXT_EXTENSIONS | METADATA_EXTENSIONS
BINARY_LIBRARY_EXTENSIONS = {".enl", ".enlx", ".enlp"}
GENERATED_SKIP_TOKENS = {
    "__pycache__", "local_first_literature_radar", "literature_gap_radar", "frontier_scan",
    "frontier_screened", "frontier_enriched", "manifest", "证据矩阵",
    "阅读优先级", "归纳说明", "文献深度清点", "全本文献深度清点",
    "已有文献总清点", "literature_evidence_matrix", "reading_priority_report",
    "literature_gap_radar",
}

STOPWORDS = {
    "the", "and", "for", "with", "from", "into", "that", "this", "these",
    "those", "using", "used", "study", "paper", "results", "result",
    "method", "methods", "effect", "effects", "based", "between", "within",
    "without", "through", "under", "over", "after", "before", "were", "was",
    "are", "is", "can", "may", "not", "have", "has", "had", "their", "our",
    "your", "which", "also", "than", "then", "such", "toward", "towards",
    "via", "new", "high", "low", "role", "roles", "analysis", "review",
}

MECHANISM_CUES = {
    "mechanism", "pathway", "signaling", "regulation", "regulates", "drives",
    "mediates", "axis", "feedback", "activation", "inhibition", "causal",
    "cascade", "transduction", "interaction", "binding", "phase separation",
    "mechanotransduction", "charge transfer", "heterojunction",
}

EXPERIMENT_CUES = {
    "assay", "staining", "immunofluorescence", "western blot", "qpcr", "rt-pcr",
    "elisa", "flow cytometry", "confocal", "sem", "tem", "xps", "xrd", "ftir",
    "rheology", "frap", "animal model", "in vivo", "in vitro", "control",
    "blank", "negative control", "positive control", "time point", "dose",
    "release", "degradation", "viability", "migration", "proliferation",
}

METHOD_SECTION_RE = re.compile(
    r"(?is)(materials and methods|methods|experimental section|methodology)"
    r".{0,12000}"
)
ABSTRACT_RE = re.compile(r"(?is)\babstract\b[:\s]*(.{100,4000}?)(?:\n\s*\n|introduction|keywords)")
FIGURE_RE = re.compile(r"(?is)(fig\.?\s*\d+|figure\s+\d+|scheme\s+\d+).{0,1200}")
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", re.I)
YEAR_RE = re.compile(r"\b(19[8-9]\d|20[0-4]\d)\b")


def eprint(*args: Any) -> None:
    print(*args, file=sys.stderr)


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def stable_hash(text: str) -> int:
    return int(hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()[:8], 16)


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def normalize_doi(doi: str) -> str:
    doi = (doi or "").strip()
    doi = re.sub(r"^https?://(dx\.)?doi\.org/", "", doi, flags=re.I)
    doi = doi.strip(" .;,\n\t")
    return doi.lower()


def normalize_title(title: str) -> str:
    title = (title or "").lower()
    title = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", " ", title)
    return normalize_space(title)


def title_from_filename(path: Path) -> str:
    stem = path.stem
    stem = re.sub(r"^\s*[\[\(]?\d{1,3}[\]\)]?[_\-\s]+", "", stem)
    stem = re.sub(r"\b(19[8-9]\d|20[0-4]\d)\b", "", stem)
    stem = re.sub(r"\bIF\s*(about|约|=)?\s*[\d.]+", "", stem, flags=re.I)
    stem = re.sub(r"[_\-]+", " ", stem)
    return normalize_space(stem)


def file_fingerprint(path: Path, chunk_size: int = 1024 * 1024) -> str:
    h = hashlib.sha1()
    with path.open("rb") as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def read_text_plain(path: Path, max_chars: int) -> Tuple[str, str]:
    for enc in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            text = path.read_text(encoding=enc, errors="ignore")
            return text[:max_chars], f"text:{enc}"
        except Exception:
            continue
    return "", "text:failed"


def read_docx(path: Path, max_chars: int) -> Tuple[str, str]:
    try:
        with zipfile.ZipFile(path) as zf:
            names = [n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml")]
            texts: List[str] = []
            for name in names:
                if not (name.endswith("document.xml") or "header" in name or "footer" in name):
                    continue
                root = ET.fromstring(zf.read(name))
                for node in root.iter():
                    if node.tag.endswith("}t") and node.text:
                        texts.append(node.text)
            return "\n".join(texts)[:max_chars], "docx:xml"
    except Exception as exc:
        return "", f"docx:failed:{type(exc).__name__}"


def read_pdf(path: Path, max_pages: int, max_chars: int) -> Tuple[str, str]:
    readers = []
    try:
        import pypdf  # type: ignore

        readers.append(("pypdf", pypdf.PdfReader))
    except Exception:
        pass
    try:
        import PyPDF2  # type: ignore

        readers.append(("PyPDF2", PyPDF2.PdfReader))
    except Exception:
        pass

    for label, reader_cls in readers:
        try:
            reader = reader_cls(str(path))
            pages = reader.pages[:max_pages]
            parts = []
            for page in pages:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    continue
            text = "\n".join(parts)
            return text[:max_chars], f"pdf:{label}:pages={len(pages)}"
        except Exception as exc:
            last_error = f"{label}:{type(exc).__name__}"
    if readers:
        return "", f"pdf:failed:{last_error}"
    return "", "pdf:no_optional_reader"


def extract_text(path: Path, max_pages: int, max_chars: int) -> Tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path, max_pages=max_pages, max_chars=max_chars)
    if suffix == ".docx":
        return read_docx(path, max_chars=max_chars)
    if suffix in TEXT_EXTENSIONS:
        return read_text_plain(path, max_chars=max_chars)
    return "", "unsupported"


def first_match(pattern: re.Pattern[str], text: str) -> str:
    m = pattern.search(text or "")
    return normalize_space(m.group(1) if m and m.groups() else m.group(0) if m else "")


def find_doi(text: str, filename: str = "") -> str:
    for source in (text, filename):
        m = DOI_RE.search(source or "")
        if m:
            return normalize_doi(m.group(0))
    return ""


def find_year(text: str, filename: str = "") -> Optional[int]:
    years = []
    for source in (filename, text[:5000] if text else ""):
        years.extend(int(m.group(1)) for m in YEAR_RE.finditer(source or ""))
    if not years:
        return None
    current = dt.date.today().year + 1
    valid = [y for y in years if 1980 <= y <= current]
    return max(valid) if valid else None


def extract_snippets(text: str) -> Dict[str, str]:
    abstract = first_match(ABSTRACT_RE, text)
    method = first_match(METHOD_SECTION_RE, text)
    figure = first_match(FIGURE_RE, text)
    return {
        "abstract_snippet": abstract[:1800],
        "methods_snippet": method[:2200],
        "figure_snippet": figure[:1600],
    }


def evidence_flags(text: str, status: str) -> Dict[str, Any]:
    lower = (text or "").lower()
    return {
        "extraction_status": status,
        "text_chars": len(text or ""),
        "has_abstract": bool(ABSTRACT_RE.search(text or "")),
        "has_methods": bool(METHOD_SECTION_RE.search(text or "")) or any(cue in lower for cue in EXPERIMENT_CUES),
        "has_figure_legend": bool(FIGURE_RE.search(text or "")),
        "manual_verification_needed": (not text) or "failed" in status or "no_optional_reader" in status,
    }


def root_for_path(path: Path, roots: Sequence[str]) -> str:
    for root in roots:
        try:
            path.relative_to(Path(root))
            return str(Path(root))
        except Exception:
            continue
    return ""


def detect_source_kind(path: Path) -> str:
    if path.is_file():
        suffix = path.suffix.lower()
        if path.name.lower() == "zotero.sqlite":
            return "zotero_sqlite"
        if suffix in {".ris", ".nbib"}:
            return "ris_export"
        if suffix == ".bib":
            return "bibtex_export"
        if suffix == ".xml":
            return "xml_export"
        if suffix in BINARY_LIBRARY_EXTENSIONS:
            return "endnote_binary"
        if suffix == ".md":
            return "obsidian_or_markdown_note"
        return "file"
    if (path / "zotero.sqlite").exists():
        return "zotero_data_directory"
    if (path / ".obsidian").exists():
        return "obsidian_vault"
    if any(p.suffix.lower() in BINARY_LIBRARY_EXTENSIONS for p in path.glob("*")):
        return "endnote_library_directory"
    return "folder"


def build_source_manifest(roots: Sequence[str]) -> Dict[str, Any]:
    entries: List[Dict[str, Any]] = []
    for root in roots:
        path = Path(root)
        entry: Dict[str, Any] = {
            "path": str(path),
            "exists": path.exists(),
            "source_kind": detect_source_kind(path) if path.exists() else "missing",
            "notes": [],
        }
        if not path.exists():
            entry["notes"].append("Path was not found and was skipped.")
        elif entry["source_kind"] in {"endnote_binary", "endnote_library_directory"}:
            entry["notes"].append("EndNote .enl/.enlx is proprietary; export RIS, BibTeX, or XML for record-level metadata. Attached PDFs are still indexed if they are under the supplied path.")
        elif entry["source_kind"] in {"zotero_sqlite", "zotero_data_directory"}:
            entry["notes"].append("Zotero metadata is read from zotero.sqlite when accessible; storage PDFs are also indexed as local full text.")
        elif entry["source_kind"] == "obsidian_vault":
            entry["notes"].append("Markdown notes are indexed as local text; linked PDFs are indexed when they live under the supplied path.")
        entries.append(entry)
    return {"created_at": dt.datetime.now().isoformat(timespec="seconds"), "sources": entries}


def record_from_metadata(
    *,
    title: str,
    abstract: str = "",
    doi: str = "",
    year: Optional[int] = None,
    journal: str = "",
    authors: str = "",
    source_path: str = "",
    root: str = "",
    source_adapter: str = "metadata",
    local_record_id: str = "",
    attachment_paths: Optional[Sequence[str]] = None,
) -> Dict[str, Any]:
    text = "\n".join(part for part in (title, abstract, journal, authors) if part)
    snippets = extract_snippets(text)
    if abstract and not snippets.get("abstract_snippet"):
        snippets["abstract_snippet"] = abstract[:1800]
    status = f"metadata:{source_adapter}"
    return {
        "source_path": source_path,
        "root": root,
        "filename": Path(source_path).name if source_path else "",
        "file_ext": Path(source_path).suffix.lower() if source_path else "",
        "file_size": Path(source_path).stat().st_size if source_path and Path(source_path).exists() else 0,
        "sha1": hashlib.sha1((source_adapter + local_record_id + title + doi).encode("utf-8", errors="ignore")).hexdigest(),
        "doi": normalize_doi(doi) or find_doi(text),
        "title": normalize_space(title),
        "title_normalized": normalize_title(title),
        "year": year or find_year(text),
        "journal": normalize_space(journal),
        "authors": normalize_space(authors),
        "source_adapter": source_adapter,
        "local_record_id": local_record_id,
        "attachment_paths": list(attachment_paths or []),
        **snippets,
        **evidence_flags(text, status),
        "manual_verification_needed": True,
    }


def parse_ris_records(path: Path, roots: Sequence[str]) -> List[Dict[str, Any]]:
    text, _ = read_text_plain(path, max_chars=20_000_000)
    raw_records = re.split(r"(?im)^\s*ER\s*-\s*$", text)
    records: List[Dict[str, Any]] = []
    for idx, raw in enumerate(raw_records, 1):
        fields: Dict[str, List[str]] = {}
        current = ""
        for line in raw.splitlines():
            m = re.match(r"^([A-Z0-9]{2})\s*-\s*(.*)$", line.strip())
            if m:
                current = m.group(1)
                fields.setdefault(current, []).append(m.group(2).strip())
            elif current and line.strip():
                fields[current][-1] += " " + line.strip()
        title = first_field(fields, ("TI", "T1", "CT", "BT"))
        if not title:
            continue
        year = parse_year(first_field(fields, ("PY", "Y1", "DA")))
        authors = "; ".join(fields.get("AU", [])[:12])
        records.append(
            record_from_metadata(
                title=title,
                abstract=first_field(fields, ("AB", "N2")),
                doi=first_field(fields, ("DO",)),
                year=year,
                journal=first_field(fields, ("JO", "JF", "JA", "T2")),
                authors=authors,
                source_path=str(path),
                root=root_for_path(path, roots),
                source_adapter="ris",
                local_record_id=f"{path.name}#{idx}",
                attachment_paths=fields.get("L1", []) + fields.get("L2", []),
            )
        )
    return records


def parse_bibtex_records(path: Path, roots: Sequence[str]) -> List[Dict[str, Any]]:
    text, _ = read_text_plain(path, max_chars=20_000_000)
    records: List[Dict[str, Any]] = []
    for idx, entry in enumerate(split_bibtex_entries(text), 1):
        fields = parse_bibtex_fields(entry)
        title = cleanup_bibtex_value(fields.get("title", ""))
        if not title:
            continue
        records.append(
            record_from_metadata(
                title=title,
                abstract=cleanup_bibtex_value(fields.get("abstract", "")),
                doi=cleanup_bibtex_value(fields.get("doi", "")),
                year=parse_year(cleanup_bibtex_value(fields.get("year", "") or fields.get("date", ""))),
                journal=cleanup_bibtex_value(fields.get("journal", "") or fields.get("journaltitle", "") or fields.get("booktitle", "")),
                authors=cleanup_bibtex_value(fields.get("author", "")),
                source_path=str(path),
                root=root_for_path(path, roots),
                source_adapter="bibtex",
                local_record_id=f"{path.name}#{idx}",
                attachment_paths=[cleanup_bibtex_value(fields.get(k, "")) for k in ("file", "url") if fields.get(k)],
            )
        )
    return records


def split_bibtex_entries(text: str) -> List[str]:
    starts = [m.start() for m in re.finditer(r"@\w+\s*\{", text)]
    entries = []
    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)
        entries.append(text[start:end])
    return entries


def parse_bibtex_fields(entry: str) -> Dict[str, str]:
    fields: Dict[str, str] = {}
    for m in re.finditer(r"(?is)\b([A-Za-z][A-Za-z0-9_-]*)\s*=\s*(\{(?:[^{}]|\{[^{}]*\})*\}|\"[^\"]*\"|[^,\n]+)", entry):
        fields[m.group(1).lower()] = m.group(2).strip().rstrip(",")
    return fields


def cleanup_bibtex_value(value: str) -> str:
    value = (value or "").strip().strip(",")
    if (value.startswith("{") and value.endswith("}")) or (value.startswith('"') and value.endswith('"')):
        value = value[1:-1]
    value = value.replace("{", "").replace("}", "")
    return normalize_space(value)


def parse_endnote_xml_records(path: Path, roots: Sequence[str]) -> List[Dict[str, Any]]:
    try:
        root = ET.parse(path).getroot()
    except Exception as exc:
        eprint(f"[warn] XML metadata parse failed: {path} ({type(exc).__name__})")
        return []
    records: List[Dict[str, Any]] = []
    for idx, rec in enumerate([n for n in root.iter() if strip_ns(n.tag).lower() == "record"], 1):
        title = xml_text_first(rec, ("title",))
        if not title:
            continue
        authors = "; ".join(xml_texts(rec, ("author",))[:12])
        records.append(
            record_from_metadata(
                title=title,
                abstract=xml_text_first(rec, ("abstract",)),
                doi=xml_text_first(rec, ("electronic-resource-num", "doi")),
                year=parse_year(xml_text_first(rec, ("year", "dates"))),
                journal=xml_text_first(rec, ("full-title", "secondary-title", "journal")),
                authors=authors,
                source_path=str(path),
                root=root_for_path(path, roots),
                source_adapter="endnote_xml",
                local_record_id=f"{path.name}#{idx}",
            )
        )
    return records


def parse_zotero_sqlite(path: Path, roots: Sequence[str]) -> List[Dict[str, Any]]:
    try:
        conn = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
        conn.row_factory = sqlite3.Row
    except Exception as exc:
        eprint(f"[warn] Zotero sqlite could not be opened: {path} ({type(exc).__name__})")
        return []
    try:
        fields_by_item: Dict[int, Dict[str, str]] = {}
        query = """
            SELECT i.itemID, i.key, it.typeName, f.fieldName, v.value
            FROM items i
            JOIN itemTypes it ON i.itemTypeID = it.itemTypeID
            JOIN itemData d ON i.itemID = d.itemID
            JOIN fields f ON d.fieldID = f.fieldID
            JOIN itemDataValues v ON d.valueID = v.valueID
            WHERE it.typeName IN ('journalArticle','conferencePaper','bookSection','preprint','thesis','report','book')
        """
        item_meta: Dict[int, Dict[str, str]] = {}
        for row in conn.execute(query):
            item_id = int(row["itemID"])
            fields_by_item.setdefault(item_id, {})[str(row["fieldName"])] = str(row["value"])
            item_meta.setdefault(item_id, {"key": str(row["key"]), "type": str(row["typeName"])})

        authors_by_item: Dict[int, List[str]] = {}
        try:
            for row in conn.execute(
                """
                SELECT ic.itemID, cd.firstName, cd.lastName, cd.fieldMode
                FROM itemCreators ic
                JOIN creators c ON ic.creatorID = c.creatorID
                JOIN creatorData cd ON c.creatorDataID = cd.creatorDataID
                ORDER BY ic.itemID, ic.orderIndex
                """
            ):
                item_id = int(row["itemID"])
                if int(row["fieldMode"] or 0) == 1:
                    name = str(row["lastName"] or "")
                else:
                    name = normalize_space(f"{row['firstName'] or ''} {row['lastName'] or ''}")
                if name:
                    authors_by_item.setdefault(item_id, []).append(name)
        except Exception:
            pass

        attachments_by_item: Dict[int, List[str]] = {}
        try:
            for row in conn.execute(
                """
                SELECT ia.parentItemID, ai.key, ia.path
                FROM itemAttachments ia
                JOIN items ai ON ia.itemID = ai.itemID
                WHERE ia.parentItemID IS NOT NULL
                """
            ):
                parent = int(row["parentItemID"])
                att_path = str(row["path"] or "")
                if att_path.startswith("storage:"):
                    att_path = str(path.parent / "storage" / str(row["key"]) / att_path.split(":", 1)[1])
                attachments_by_item.setdefault(parent, []).append(att_path)
        except Exception:
            pass

        records = []
        for item_id, fields in fields_by_item.items():
            title = fields.get("title", "")
            if not title:
                continue
            records.append(
                record_from_metadata(
                    title=title,
                    abstract=fields.get("abstractNote", ""),
                    doi=fields.get("DOI", ""),
                    year=parse_year(fields.get("date", "")),
                    journal=fields.get("publicationTitle", "") or fields.get("journalAbbreviation", ""),
                    authors="; ".join(authors_by_item.get(item_id, [])[:12]),
                    source_path=str(path),
                    root=root_for_path(path, roots),
                    source_adapter="zotero_sqlite",
                    local_record_id=f"zotero:{item_meta.get(item_id, {}).get('key', item_id)}",
                    attachment_paths=attachments_by_item.get(item_id, []),
                )
            )
        return records
    except Exception as exc:
        eprint(f"[warn] Zotero sqlite parse failed: {path} ({type(exc).__name__})")
        return []
    finally:
        conn.close()


def strip_ns(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def xml_texts(node: ET.Element, names: Sequence[str]) -> List[str]:
    wanted = {n.lower() for n in names}
    values = []
    for child in node.iter():
        if strip_ns(child.tag).lower() in wanted:
            text = normalize_space(" ".join(t.strip() for t in child.itertext() if t and t.strip()))
            if text:
                values.append(text)
    return values


def xml_text_first(node: ET.Element, names: Sequence[str]) -> str:
    values = xml_texts(node, names)
    return values[0] if values else ""


def first_field(fields: Dict[str, List[str]], names: Sequence[str]) -> str:
    for name in names:
        values = fields.get(name)
        if values:
            return normalize_space(values[0])
    return ""


def parse_year(value: str) -> Optional[int]:
    m = YEAR_RE.search(value or "")
    return int(m.group(1)) if m else None


def iter_metadata_records(roots: Sequence[str]) -> Iterable[Dict[str, Any]]:
    seen_files: set[str] = set()
    for root in roots:
        root_path = Path(root)
        if not root_path.exists():
            continue
        candidates: List[Path] = []
        if root_path.is_file():
            candidates.append(root_path)
        else:
            zotero = root_path / "zotero.sqlite"
            if zotero.exists():
                candidates.append(zotero)
            candidates.extend(p for p in root_path.rglob("*") if p.is_file() and p.suffix.lower() in METADATA_EXTENSIONS)
        for path in candidates:
            key = str(path.resolve()).lower()
            if key in seen_files or should_skip_index_file(path):
                continue
            seen_files.add(key)
            suffix = path.suffix.lower()
            if path.name.lower() == "zotero.sqlite":
                yield from parse_zotero_sqlite(path, roots)
            elif suffix in {".ris", ".nbib"}:
                yield from parse_ris_records(path, roots)
            elif suffix == ".bib":
                yield from parse_bibtex_records(path, roots)
            elif suffix == ".xml":
                yield from parse_endnote_xml_records(path, roots)


def iter_literature_files(roots: Sequence[str]) -> Iterable[Path]:
    seen = set()
    for root in roots:
        root_path = Path(root)
        if not root_path.exists():
            eprint(f"[warn] root not found: {root}")
            continue
        files = [root_path] if root_path.is_file() else root_path.rglob("*")
        for path in files:
            if not path.is_file():
                continue
            if path.suffix.lower() not in FULLTEXT_EXTENSIONS:
                continue
            if should_skip_index_file(path):
                continue
            key = str(path.resolve()).lower()
            if key in seen:
                continue
            seen.add(key)
            yield path


def should_skip_index_file(path: Path) -> bool:
    low = str(path).lower()
    name = path.name.lower()
    if name.startswith("~$"):
        return True
    if any(token.lower() in low for token in GENERATED_SKIP_TOKENS):
        return True
    if path.suffix.lower() in {".docx", ".md", ".txt"} and (name.startswith("cx_") or name.startswith("cl_")):
        return True
    return False


def build_local_record(path: Path, roots: Sequence[str], max_pages: int, max_chars: int) -> Dict[str, Any]:
    text, status = extract_text(path, max_pages=max_pages, max_chars=max_chars)
    snippets = extract_snippets(text)
    doi = find_doi(text, path.name)
    year = find_year(text, path.name)
    title = title_from_filename(path)
    root_label = root_for_path(path, roots)
    source_kind = "obsidian_markdown" if path.suffix.lower() == ".md" and any(part.lower() == ".obsidian" for part in path.parts) else "file"
    return {
        "source_path": str(path),
        "root": root_label,
        "filename": path.name,
        "file_ext": path.suffix.lower(),
        "file_size": path.stat().st_size,
        "sha1": file_fingerprint(path),
        "doi": doi,
        "title": title,
        "title_normalized": normalize_title(title),
        "year": year,
        "source_adapter": source_kind,
        "local_record_id": str(path),
        "attachment_paths": [],
        **snippets,
        **evidence_flags(text, status),
    }


def write_jsonl(path: Path, records: Iterable[Dict[str, Any]]) -> int:
    ensure_parent(path)
    count = 0
    with path.open("w", encoding="utf-8", newline="\n") as f:
        for rec in records:
            f.write(json.dumps(rec, ensure_ascii=False, sort_keys=True) + "\n")
            count += 1
    return count


def dedupe_local_records(records: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    kept: List[Dict[str, Any]] = []
    seen_doi: set[str] = set()
    seen_title: set[str] = set()
    for rec in records:
        doi = normalize_doi(rec.get("doi", ""))
        title = rec.get("title_normalized") or normalize_title(rec.get("title", ""))
        if doi and doi in seen_doi:
            continue
        if not doi and title and title in seen_title:
            continue
        if doi:
            seen_doi.add(doi)
        if title:
            seen_title.add(title)
        kept.append(rec)
    return kept


def write_source_manifest(path: Optional[Path], roots: Sequence[str], indexed_count: int) -> None:
    if not path:
        return
    manifest = build_source_manifest(roots)
    manifest["indexed_record_count"] = indexed_count
    ensure_parent(path)
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    eprint(f"[ok] wrote source manifest: {path}")


def read_json_any(path: Path) -> Any:
    text = path.read_text(encoding="utf-8-sig", errors="ignore").strip()
    if not text:
        return []
    if text[0] in "[{":
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
    return [json.loads(line) for line in text.splitlines() if line.strip()]


def read_records(path: Path) -> List[Dict[str, Any]]:
    data = read_json_any(path)
    if isinstance(data, dict):
        if "records" in data and isinstance(data["records"], list):
            return data["records"]
        if "new" in data and isinstance(data["new"], list):
            return data["new"]
        return [data]
    return list(data)


def simple_yaml_like(path: Path) -> Dict[str, Any]:
    """Parse a small top-level YAML subset without PyYAML."""
    result: Dict[str, Any] = {}
    current_key: Optional[str] = None
    for raw in path.read_text(encoding="utf-8-sig", errors="ignore").splitlines():
        line = raw.split("#", 1)[0].rstrip()
        if not line:
            continue
        if not raw.startswith(" ") and ":" in line:
            key, value = line.split(":", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if value:
                result[key] = value
                current_key = None
            else:
                result[key] = []
                current_key = key
            continue
        if current_key and line.strip().startswith("-"):
            result.setdefault(current_key, []).append(line.strip()[1:].strip().strip('"').strip("'"))
    return result


def load_profile(path: Optional[Path]) -> Dict[str, Any]:
    if not path:
        return {}
    if path.suffix.lower() == ".json":
        data = read_json_any(path)
        return data if isinstance(data, dict) else {}
    return simple_yaml_like(path)


def term_tokens(text: str) -> List[str]:
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9+\-/]{2,}|[\u4e00-\u9fff]{2,}", text or "")
    cleaned = []
    for token in tokens:
        t = token.strip("-_/").lower()
        if len(t) < 3 or t in STOPWORDS:
            continue
        if t.isdigit():
            continue
        cleaned.append(t)
    return cleaned


def top_terms(records: Sequence[Dict[str, Any]], topn: int = 40) -> List[str]:
    counter: Counter[str] = Counter()
    for rec in records:
        text = " ".join(
            str(rec.get(k, ""))
            for k in ("title", "abstract_snippet", "methods_snippet", "figure_snippet")
        )
        counter.update(term_tokens(text))
    terms = [term for term, _ in counter.most_common(topn * 3)]
    filtered = []
    for term in terms:
        if any(term in kept or kept in term for kept in filtered):
            continue
        filtered.append(term)
        if len(filtered) >= topn:
            break
    return filtered


def infer_profile_from_index(index_path: Path, output: Path, design_paths: Sequence[str], topn: int) -> None:
    records = read_records(index_path)
    design_text = read_many_design_texts(design_paths)
    terms = top_terms(records, topn=topn)
    design_terms = [t for t in term_tokens(design_text) if t not in STOPWORDS]
    design_counts = Counter(design_terms)
    method_terms = [t for t in terms if any(cue in t for cue in ("pcr", "assay", "staining", "seq", "rheology", "microscopy", "xps", "sem", "tem"))]
    profile = {
        "profile_name": "inferred-literature-gap-radar-profile",
        "created_at": dt.datetime.now().isoformat(timespec="seconds"),
        "inferred_from_index": str(index_path),
        "local_record_count": len(records),
        "core_terms": terms[:18],
        "secondary_terms": terms[18:36],
        "method_terms": method_terms[:16],
        "design_terms": [term for term, _ in design_counts.most_common(24)],
        "exclusion_topics": [],
        "must_track_journals": [],
        "nice_to_track_journals": [],
        "weekly_reading_capacity": 5,
        "scoring_weights": DEFAULT_WEIGHTS,
    }
    ensure_parent(output)
    output.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    eprint(f"[ok] wrote profile: {output}")


def read_many_design_texts(paths: Sequence[str]) -> str:
    parts: List[str] = []
    for item in paths or []:
        path = Path(item)
        if not path.exists():
            eprint(f"[warn] design path not found: {item}")
            continue
        files = [path] if path.is_file() else [p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in DOC_EXTENSIONS]
        for f in files[:80]:
            text, status = extract_text(f, max_pages=80, max_chars=160000)
            if text:
                parts.append(f"\n\n# SOURCE: {f}\n{text}")
            else:
                eprint(f"[warn] could not read design/source text: {f} ({status})")
    return "\n".join(parts)


def build_gap_map(profile_path: Path, output: Path, design_paths: Sequence[str]) -> None:
    profile = load_profile(profile_path)
    design_text = read_many_design_texts(design_paths)
    terms = []
    for key in ("core_terms", "core_topics", "secondary_terms", "secondary_topics", "mechanism_terms", "method_terms"):
        value = profile.get(key, [])
        if isinstance(value, str):
            terms.append(value)
        else:
            terms.extend(value)
    terms = [normalize_space(str(t)) for t in terms if normalize_space(str(t))]
    if design_text:
        design_counter = Counter(term_tokens(design_text))
        design_terms = [term for term, _ in design_counter.most_common(30)]
    else:
        design_terms = []
    base_terms = (terms[:18] + design_terms[:12])[:24]

    gap_specs = [
        ("background_context", "synthesis", ["review", "overview", "state of the art"]),
        ("mechanism_support", "mechanism", ["mechanism", "pathway", "signaling", "causal"]),
        ("experimental_methods", "method", ["assay", "protocol", "quantification", "control"]),
        ("controls_and_benchmarks", "control", ["control", "benchmark", "comparison", "standard"]),
        ("competing_or_adjacent_systems", "benchmark", ["similar system", "alternative approach", "comparative"]),
        ("figures_and_writing", "writing", ["figure", "schematic", "graphical abstract", "workflow"]),
    ]
    gaps = []
    for name, need_type, cues in gap_specs:
        query_terms = list(dict.fromkeys(base_terms[:8] + cues))
        evidence = find_sentences_for_cues(design_text, cues, limit=3) if design_text else []
        gaps.append(
            {
                "gap_id": name,
                "need_type": need_type,
                "question": default_gap_question(name),
                "query_terms": query_terms,
                "design_evidence": evidence,
            }
        )
    ensure_parent(output)
    output.write_text(json.dumps({"created_at": dt.datetime.now().isoformat(timespec="seconds"), "gaps": gaps}, ensure_ascii=False, indent=2), encoding="utf-8")
    eprint(f"[ok] wrote gap map: {output}")


def default_gap_question(name: str) -> str:
    questions = {
        "background_context": "Which recent papers best establish the field context and unmet need?",
        "mechanism_support": "Which papers support or challenge the mechanism claimed by the project?",
        "experimental_methods": "Which papers provide reusable experimental design, controls, parameters, or readouts?",
        "controls_and_benchmarks": "Which papers define appropriate negative controls, positive controls, or benchmark systems?",
        "competing_or_adjacent_systems": "Which adjacent systems or competing approaches should be cited or compared?",
        "figures_and_writing": "Which papers provide figure logic, graphical abstract structure, or evidence-chain presentation?",
    }
    return questions.get(name, name)


def find_sentences_for_cues(text: str, cues: Sequence[str], limit: int = 3) -> List[str]:
    if not text:
        return []
    sentences = re.split(r"(?<=[.!?。！？])\s+", normalize_space(text))
    found = []
    for sent in sentences:
        lower = sent.lower()
        if any(cue.lower() in lower for cue in cues):
            found.append(sent[:360])
            if len(found) >= limit:
                break
    return found


def openalex_abstract(inv: Optional[Dict[str, List[int]]]) -> str:
    if not inv:
        return ""
    max_pos = max((max(v) for v in inv.values() if v), default=-1)
    words = [""] * (max_pos + 1)
    for word, positions in inv.items():
        for pos in positions:
            if 0 <= pos <= max_pos:
                words[pos] = word
    return normalize_space(" ".join(words))


def candidate_from_openalex(work: Dict[str, Any], query: str) -> Dict[str, Any]:
    source = ((work.get("primary_location") or {}).get("source") or {}) or (work.get("host_venue") or {})
    authorships = work.get("authorships") or []
    authors = []
    for a in authorships[:8]:
        author = a.get("author") or {}
        if author.get("display_name"):
            authors.append(author["display_name"])
    return {
        "source": "OpenAlex",
        "query": query,
        "openalex_id": work.get("id", ""),
        "doi": normalize_doi(work.get("doi") or ""),
        "title": normalize_space(work.get("title") or ""),
        "year": work.get("publication_year"),
        "publication_date": work.get("publication_date"),
        "journal": source.get("display_name", ""),
        "authors": "; ".join(authors),
        "abstract": openalex_abstract(work.get("abstract_inverted_index")),
        "url": work.get("doi") or work.get("id") or "",
        "cited_by_count": work.get("cited_by_count"),
        "type": work.get("type"),
    }


def load_gap_map(path: Optional[Path]) -> Dict[str, Any]:
    if not path or not path.exists():
        return {"gaps": []}
    data = read_json_any(path)
    return data if isinstance(data, dict) else {"gaps": data}


def build_queries(profile: Dict[str, Any], gap_map: Dict[str, Any], max_queries: int) -> List[str]:
    queries: List[str] = []
    for gap in gap_map.get("gaps", []):
        terms = gap.get("query_terms", [])
        if terms:
            queries.append(" ".join(str(t) for t in terms[:6]))
    core = []
    for key in ("core_terms", "core_topics", "design_terms", "method_terms"):
        value = profile.get(key, [])
        if isinstance(value, str):
            core.append(value)
        else:
            core.extend(value)
    core = [str(t) for t in core if str(t).strip()]
    if core:
        queries.append(" ".join(core[:6]))
    if len(core) >= 6:
        queries.append(" ".join(core[3:9]))
    seen = set()
    unique = []
    for q in queries:
        q = normalize_space(q)
        if len(q) < 4:
            continue
        key = q.lower()
        if key not in seen:
            unique.append(q)
            seen.add(key)
    return unique[:max_queries]


def search_openalex(profile_path: Path, gap_map_path: Optional[Path], output: Path, years: int, from_date: Optional[str], to_date: Optional[str], per_query: int, max_queries: int, mailto: str) -> None:
    profile = load_profile(profile_path)
    gap_map = load_gap_map(gap_map_path)
    today = dt.date.today()
    end = dt.date.fromisoformat(to_date) if to_date else today
    start = dt.date.fromisoformat(from_date) if from_date else dt.date(end.year - years, end.month, end.day)
    queries = build_queries(profile, gap_map, max_queries=max_queries)
    if not queries:
        raise SystemExit("No search queries could be built from profile/gap map.")

    seen = set()
    records: List[Dict[str, Any]] = []
    for query in queries:
        params = {
            "search": query,
            "filter": f"from_publication_date:{start.isoformat()},to_publication_date:{end.isoformat()},type:article",
            "per-page": str(per_query),
            "sort": "publication_date:desc",
        }
        if mailto:
            params["mailto"] = mailto
        url = "https://api.openalex.org/works?" + urllib.parse.urlencode(params)
        eprint(f"[search] {query}")
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "literature-gap-radar/0.2"})
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read().decode("utf-8"))
        except Exception as exc:
            eprint(f"[warn] OpenAlex query failed: {type(exc).__name__}: {exc}")
            continue
        for work in data.get("results", []):
            rec = candidate_from_openalex(work, query=query)
            key = rec.get("doi") or rec.get("openalex_id") or normalize_title(rec.get("title", ""))
            if key in seen:
                continue
            seen.add(key)
            records.append(rec)
        time.sleep(0.15)
    count = write_jsonl(output, records)
    eprint(f"[ok] wrote {count} candidates: {output}")


def similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


def term_hit_score(terms: Sequence[str], text: str) -> float:
    if not terms or not text:
        return 0.0
    lower = text.lower()
    hits = 0
    weighted = 0.0
    for term in terms:
        t = str(term).strip().lower()
        if len(t) < 3:
            continue
        if t in lower:
            hits += 1
            weighted += 1.5 if " " in t or "-" in t else 1.0
    coverage = hits / max(1, min(len(terms), 18))
    return min(100.0, weighted * 10.0 + coverage * 45.0)


def cue_score(cues: Sequence[str], text: str) -> float:
    return term_hit_score(cues, text)


def classify_duplicate(candidate: Dict[str, Any], locals_: Sequence[Dict[str, Any]]) -> Tuple[str, float, str]:
    doi = normalize_doi(candidate.get("doi", ""))
    title_norm = normalize_title(candidate.get("title", ""))
    best_sim = 0.0
    best_path = ""
    for rec in locals_:
        local_doi = normalize_doi(rec.get("doi", ""))
        if doi and local_doi and doi == local_doi:
            return "already", 1.0, rec.get("source_path", "")
        sim = similarity(title_norm, rec.get("title_normalized") or normalize_title(rec.get("title", "")))
        if sim > best_sim:
            best_sim = sim
            best_path = rec.get("source_path", "")
    if best_sim >= 0.94:
        return "already", best_sim, best_path
    if best_sim >= 0.82:
        return "near_duplicate", best_sim, best_path
    return "new", best_sim, best_path


def journal_year_score(candidate: Dict[str, Any], profile: Dict[str, Any]) -> float:
    year = candidate.get("year")
    try:
        year_i = int(year)
    except Exception:
        year_i = None
    current = dt.date.today().year
    recency = 35.0
    if year_i:
        age = max(0, current - year_i)
        recency = max(20.0, 75.0 - age * 8.0)
    journal = (candidate.get("journal") or "").lower()
    must = [str(j).lower() for j in profile.get("must_track_journals", []) or []]
    nice = [str(j).lower() for j in profile.get("nice_to_track_journals", []) or []]
    boost = 0.0
    if journal and any(j in journal or journal in j for j in must):
        boost = 25.0
    elif journal and any(j in journal or journal in j for j in nice):
        boost = 14.0
    return min(100.0, recency + boost)


def evidence_score(candidate: Dict[str, Any]) -> float:
    text = " ".join(str(candidate.get(k, "")) for k in ("title", "abstract", "abstract_snippet", "methods_snippet", "figure_snippet"))
    score = 15.0
    if candidate.get("abstract") or candidate.get("abstract_snippet"):
        score += 35.0
    if len(text) > 1200:
        score += 15.0
    lower = text.lower()
    if any(cue in lower for cue in EXPERIMENT_CUES):
        score += 15.0
    if candidate.get("methods_snippet"):
        score += 15.0
    if candidate.get("figure_snippet"):
        score += 10.0
    return min(100.0, score)


def score_candidates(index_path: Path, candidates_path: Path, profile_path: Path, gap_map_path: Optional[Path], output: Path) -> None:
    locals_ = read_records(index_path) if index_path.exists() else []
    candidates = read_records(candidates_path)
    profile = load_profile(profile_path)
    gap_map = load_gap_map(gap_map_path)
    profile_terms: List[str] = []
    for key in ("core_terms", "core_topics", "secondary_terms", "secondary_topics", "design_terms", "method_terms"):
        val = profile.get(key, [])
        profile_terms.extend([val] if isinstance(val, str) else list(val or []))
    gap_terms: List[str] = []
    for gap in gap_map.get("gaps", []):
        gap_terms.extend(gap.get("query_terms", []) or [])
    weights = dict(DEFAULT_WEIGHTS)
    if isinstance(profile.get("scoring_weights"), dict):
        weights.update({k: float(v) for k, v in profile["scoring_weights"].items() if k in weights})

    scored: List[Dict[str, Any]] = []
    for cand in candidates:
        dup_label, dup_sim, dup_path = classify_duplicate(cand, locals_)
        text = " ".join(str(cand.get(k, "")) for k in ("title", "abstract", "journal", "query"))
        topic = term_hit_score(profile_terms, text)
        gap = term_hit_score(gap_terms, text)
        mechanism = cue_score(MECHANISM_CUES, text)
        experiment = cue_score(EXPERIMENT_CUES, text)
        journal = journal_year_score(cand, profile)
        evidence = evidence_score(cand)
        if dup_label == "already":
            novelty = 18.0
        elif dup_label == "near_duplicate":
            novelty = 42.0
        else:
            novelty = 75.0 + min(15.0, gap / 7.0)
        axes = {
            "topic_relevance": round(topic, 1),
            "mechanism_value": round(mechanism, 1),
            "experiment_reusability": round(experiment, 1),
            "gap_fill": round(gap, 1),
            "novelty_against_local": round(novelty, 1),
            "journal_year_signal": round(journal, 1),
            "evidence_richness": round(evidence, 1),
        }
        total = sum(axes[k] * weights[k] for k in weights)
        if dup_label == "already":
            total -= 14
        elif dup_label == "near_duplicate":
            total -= 6
        tie_breaker = (stable_hash((cand.get("doi") or cand.get("title") or "") + str(cand.get("year"))) % 19) / 10.0
        total = max(0.0, min(100.0, total + tie_breaker))
        if dup_label == "already" and total < 55:
            final_label = "low_priority_duplicate"
        elif dup_label in {"already", "near_duplicate"}:
            final_label = dup_label
        elif gap >= 55 or experiment >= 60:
            final_label = "worth_adding"
        else:
            final_label = "genuinely_new"
        rationale = build_rationale(cand, axes, final_label, dup_path)
        scored.append(
            {
                **cand,
                "dedup_label": final_label,
                "duplicate_similarity": round(dup_sim, 3),
                "nearest_local_path": dup_path,
                **{f"score_{k}": v for k, v in axes.items()},
                "score_total": round(total, 1),
                "rationale": rationale,
            }
        )
    scored.sort(key=lambda r: (r.get("score_total", 0), r.get("publication_date", "")), reverse=True)
    count = write_jsonl(output, scored)
    eprint(f"[ok] wrote {count} scored records: {output}")


def build_rationale(cand: Dict[str, Any], axes: Dict[str, float], label: str, dup_path: str) -> str:
    strengths = sorted(axes.items(), key=lambda kv: kv[1], reverse=True)[:3]
    words = ", ".join(f"{k} {v:g}" for k, v in strengths)
    if label in {"already", "near_duplicate", "low_priority_duplicate"}:
        return f"{label}; strongest signals: {words}; compare local file: {dup_path or 'not recorded'}."
    if label == "worth_adding":
        return f"Worth adding because it is new locally and scores strongly on {words}."
    return f"Genuinely new candidate; strongest signals: {words}."


def render_outputs(scored_path: Path, output_dir: Path, top: int) -> None:
    ensure_dir(output_dir)
    records = read_records(scored_path)
    csv_path = output_dir / "literature_evidence_matrix.csv"
    html_path = output_dir / "reading_priority_report.html"
    write_csv(csv_path, records)
    write_html(html_path, records[:top])
    try_write_xlsx(output_dir / "literature_evidence_matrix.xlsx", records)
    try_write_docx(output_dir / "reading_priority_report.docx", records[:top])
    eprint(f"[ok] wrote report outputs: {output_dir}")


def write_csv(path: Path, records: Sequence[Dict[str, Any]]) -> None:
    ensure_parent(path)
    fields = [
        "score_total", "dedup_label", "title", "year", "journal", "doi", "url",
        "score_topic_relevance", "score_gap_fill", "score_experiment_reusability",
        "score_mechanism_value", "score_novelty_against_local",
        "score_journal_year_signal", "score_evidence_richness",
        "duplicate_similarity", "nearest_local_path", "rationale", "abstract",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for rec in records:
            writer.writerow(rec)


def write_html(path: Path, records: Sequence[Dict[str, Any]]) -> None:
    ensure_parent(path)
    rows = []
    for rec in records:
        rows.append(
            "<tr>"
            f"<td>{html.escape(str(rec.get('score_total', '')))}</td>"
            f"<td>{html.escape(str(rec.get('dedup_label', '')))}</td>"
            f"<td>{html.escape(str(rec.get('title', '')))}</td>"
            f"<td>{html.escape(str(rec.get('year', '')))}</td>"
            f"<td>{html.escape(str(rec.get('journal', '')))}</td>"
            f"<td>{html.escape(str(rec.get('rationale', '')))}</td>"
            "</tr>"
        )
    doc = f"""<!doctype html>
<meta charset="utf-8">
<title>Literature Gap Radar report</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 32px; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 6px; vertical-align: top; }}
th {{ background: #f3f4f6; }}
</style>
<h1>Literature Gap Radar report</h1>
<p>Generated {dt.datetime.now().isoformat(timespec='seconds')}</p>
<table>
<tr><th>Score</th><th>Status</th><th>Title</th><th>Year</th><th>Journal</th><th>Rationale</th></tr>
{''.join(rows)}
</table>
"""
    path.write_text(doc, encoding="utf-8")


def try_write_xlsx(path: Path, records: Sequence[Dict[str, Any]]) -> None:
    try:
        import openpyxl  # type: ignore
    except Exception:
        eprint("[warn] openpyxl unavailable; skipped xlsx")
        return
    fields = [
        "score_total", "dedup_label", "title", "year", "journal", "doi", "url",
        "score_topic_relevance", "score_gap_fill", "score_experiment_reusability",
        "score_mechanism_value", "score_novelty_against_local",
        "score_journal_year_signal", "score_evidence_richness", "rationale",
    ]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "evidence_matrix"
    ws.append(fields)
    for rec in records:
        ws.append([rec.get(f, "") for f in fields])
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = min(60, max(12, max(len(str(c.value or "")) for c in col[:100]) + 2))
    ensure_parent(path)
    wb.save(path)


def try_write_docx(path: Path, records: Sequence[Dict[str, Any]]) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception:
        eprint("[warn] python-docx unavailable; skipped docx")
        return
    doc = Document()
    doc.add_heading("Literature Gap Radar reading priorities", 0)
    doc.add_paragraph(f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}")
    for idx, rec in enumerate(records, 1):
        doc.add_heading(f"{idx}. [{rec.get('score_total', '')}] {rec.get('title', '')}", level=2)
        doc.add_paragraph(f"Status: {rec.get('dedup_label', '')}; Year: {rec.get('year', '')}; Journal: {rec.get('journal', '')}")
        if rec.get("doi"):
            doc.add_paragraph(f"DOI: {rec.get('doi')}")
        doc.add_paragraph(f"Why: {rec.get('rationale', '')}")
        axes = [
            f"topic {rec.get('score_topic_relevance', '')}",
            f"gap {rec.get('score_gap_fill', '')}",
            f"experiment {rec.get('score_experiment_reusability', '')}",
            f"mechanism {rec.get('score_mechanism_value', '')}",
            f"novelty {rec.get('score_novelty_against_local', '')}",
        ]
        doc.add_paragraph("Scores: " + "; ".join(axes))
    ensure_parent(path)
    doc.save(path)


def command_index(args: argparse.Namespace) -> None:
    metadata_records = list(iter_metadata_records(args.roots))
    file_records = [
        build_local_record(path, args.roots, max_pages=args.max_pages, max_chars=args.max_chars)
        for path in iter_literature_files(args.roots)
    ]
    records = dedupe_local_records([*metadata_records, *file_records])
    count = write_jsonl(Path(args.output), records)
    write_source_manifest(Path(args.source_manifest) if getattr(args, "source_manifest", None) else None, args.roots, count)
    eprint(f"[ok] indexed {count} local literature records: {args.output}")


def command_run(args: argparse.Namespace) -> None:
    out = Path(args.output_dir)
    ensure_dir(out)
    index = out / "local_library_index.jsonl"
    profile = out / "inferred_profile.json"
    gap_map = out / "gap_map.json"
    candidates = out / "candidates.jsonl"
    scored = out / "scored.jsonl"
    command_index(argparse.Namespace(roots=args.roots, output=str(index), source_manifest=str(out / "source_manifest.json"), max_pages=args.max_pages, max_chars=args.max_chars))
    infer_profile_from_index(index, profile, args.design or [], topn=48)
    build_gap_map(profile, gap_map, args.design or [])
    search_openalex(profile, gap_map, candidates, years=args.years, from_date=args.from_date, to_date=args.to_date, per_query=args.per_query, max_queries=args.max_queries, mailto=args.mailto)
    score_candidates(index, candidates, profile, gap_map, scored)
    render_outputs(scored, out / "report", top=args.top)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Literature Gap Radar")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("index", help="Index local literature sources and files")
    p.add_argument("--roots", nargs="+", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--source-manifest")
    p.add_argument("--max-pages", type=int, default=80)
    p.add_argument("--max-chars", type=int, default=260000)
    p.set_defaults(func=command_index)

    p = sub.add_parser("infer-profile", help="Infer a research profile from a local index")
    p.add_argument("--index", required=True)
    p.add_argument("--design", nargs="*", default=[])
    p.add_argument("--output", required=True)
    p.add_argument("--topn", type=int, default=48)
    p.set_defaults(func=lambda a: infer_profile_from_index(Path(a.index), Path(a.output), a.design, a.topn))

    p = sub.add_parser("gap-map", help="Build a gap map from profile and optional design files")
    p.add_argument("--profile", required=True)
    p.add_argument("--design", nargs="*", default=[])
    p.add_argument("--output", required=True)
    p.set_defaults(func=lambda a: build_gap_map(Path(a.profile), Path(a.output), a.design))

    p = sub.add_parser("search", help="Search OpenAlex after profile/gap inference")
    p.add_argument("--profile", required=True)
    p.add_argument("--gap-map")
    p.add_argument("--years", type=int, default=3)
    p.add_argument("--from-date")
    p.add_argument("--to-date")
    p.add_argument("--per-query", type=int, default=25)
    p.add_argument("--max-queries", type=int, default=8)
    p.add_argument("--mailto", default="")
    p.add_argument("--output", required=True)
    p.set_defaults(func=lambda a: search_openalex(Path(a.profile), Path(a.gap_map) if a.gap_map else None, Path(a.output), a.years, a.from_date, a.to_date, a.per_query, a.max_queries, a.mailto))

    p = sub.add_parser("score", help="Deduplicate and score candidates")
    p.add_argument("--index", required=True)
    p.add_argument("--candidates", required=True)
    p.add_argument("--profile", required=True)
    p.add_argument("--gap-map")
    p.add_argument("--output", required=True)
    p.set_defaults(func=lambda a: score_candidates(Path(a.index), Path(a.candidates), Path(a.profile), Path(a.gap_map) if a.gap_map else None, Path(a.output)))

    p = sub.add_parser("render", help="Render CSV/HTML and optional XLSX/DOCX outputs")
    p.add_argument("--scored", required=True)
    p.add_argument("--output-dir", required=True)
    p.add_argument("--top", type=int, default=40)
    p.set_defaults(func=lambda a: render_outputs(Path(a.scored), Path(a.output_dir), a.top))

    p = sub.add_parser("run", help="Run the full Literature Gap Radar pipeline")
    p.add_argument("--roots", nargs="+", required=True)
    p.add_argument("--design", nargs="*", default=[])
    p.add_argument("--years", type=int, default=3)
    p.add_argument("--from-date")
    p.add_argument("--to-date")
    p.add_argument("--per-query", type=int, default=25)
    p.add_argument("--max-queries", type=int, default=8)
    p.add_argument("--mailto", default="")
    p.add_argument("--output-dir", required=True)
    p.add_argument("--max-pages", type=int, default=80)
    p.add_argument("--max-chars", type=int, default=260000)
    p.add_argument("--top", type=int, default=40)
    p.set_defaults(func=command_run)
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    args.func(args)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
